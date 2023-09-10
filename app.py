from pdf2docx import Converter
import docx
import os
from pdf2docx import parse
import openai
from metaphor_python import Metaphor
from pathlib import Path
import docx2txt

from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Get API keys from environment variables
METAPHOR_API_KEY = os.getenv("METAPHOR_API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")



# converting text from pdf to test
def getPDFtext(filename):
    docx_file = open('example.docx',"w")
    font_dir = "/usr/share/fonts"
    cv = Converter(filename)
    cv.font_dir = font_dir
    cv.convert(docx_file)      # all pages by default
    cv.close()
    #parse(filename, 'example.docx')

#converting a doc to a text file
def getDocxTextf(filename):
    doc = docx.Document(filename)
    text = open("input_text.text", "a")
    for para in doc.paragraphs:
        text.write(para.text)

#converting doc 
def getDocxText(filename):
    #doc = docx.Document(filename)
    #text = []
    #print(len(doc.parapraphs))
    #print(doc.paragraphs.text)
    #for paragraph in doc.paragraphs:
    #    text.append(paragraph.text)
    #return text
    return docx2txt.process(filename)

def getDocxText(filename):
    #doc = docx.Document(filename)
    #text = []
    #print(len(doc.parapraphs))
    #print(doc.paragraphs.text)
    #for paragraph in doc.paragraphs:
    #    text.append(paragraph.text)
    #return text
    text = docx2txt.process(filename)
    chunks = []
    for i in range(0, len(text), 2000):
        chunk = text[i:i + 2000]
        chunks.append(chunk)
    return chunks

def author_title(text, question):
    answers = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Answer the following question based on these text:\n\n{text[0]}\n\n Question: {question[0]} \n ",
        temperature=0.5,
        max_tokens=100,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )
    return answers


                        
#summary of text
def summary(text, question):
    answers = openai.Completion.create(
        engine="text-davinci-002",
        prompt=f"Answer the following questions based on these texts:\n\n{text[1:50]}\n\n Questions: {question[1:]} \n ",
        temperature=0.5,
        max_tokens=100,
        top_p=1,
        frequency_penalty=0,
        presence_penalty=0,
    )

    answer_text = open("answers.text", "w")
    answer_text.write(answers)

#find articles that this source references
def metaphor_check():
    #sources = metaphor.search(
    #"What are 3 sources that would help me understand the background of the article with the following author and title: Spoofing and countermeasures for speaker verification: a survey by Zhizheng Wua,1,∗, Nicholas Evansb, Tomi Kinnunenc, Junichi Yamagishid,e, Federico Alegreb, Haizhou Lia,fa",
    #num_results=3,
    #use_autoprompt=True,
    #type="keyword",
    #)

    search_response = metaphor.search("What are 3 sources that would help me understand the background of the article with the following author and title: Spoofing and countermeasures for speaker verification: a survey by Zhizheng Wua,1,∗, Nicholas Evansb, Tomi Kinnunenc, Junichi Yamagishid,e, Federico Alegreb, Haizhou Lia,fa",
                                      num_results = 3)
    contents_response = search_response.get_contents()

    # Print content for each result
    for content in contents_response.contents:
        print(f"Title: {content.title}\nURL: {content.url}\n")

    #find articles that reference this source
    references = metaphor.search(
    "What are 3 articles similiar to the article with the following author and title: Spoofing and countermeasures for speaker verification: a survey by Zhizheng Wua,1,∗, Nicholas Evansb, Tomi Kinnunenc, Junichi Yamagishid,e, Federico Alegreb, Haizhou Lia,f",
    num_results=3,
    use_autoprompt=True,
    type="keyword",
    )

    contents_response = references.get_contents()

    # Print content for each result
    for content in contents_response.contents:
        print(f"Title: {content.title}\nURL: {content.url}\n")


question = ["What is the title of the paper and who are the authors?", "What is the independent variable for this experiment?", "What is the dependent variable?", 
"What are top 5 most significant results?", "Why is this experiment novel?", "What model was used?",
"Explain the model architecture briefly."]

chunks = (getDocxText('example.docx'))
metaphor_check()




